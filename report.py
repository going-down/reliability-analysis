from docx import Document


DEVICE_SCHEME = 'device-scheme'
LOADS = 'loads'
FNS = 'fns'


class DumpAble:
    def dump(self):
        pass


def title(doc: Document, data):
    doc.add_paragraph("Міністерство освіти та науки, молоді та спорту України")
    doc.add_paragraph("Національний технічний університет України «КПІ»")
    doc.add_paragraph("Факультет прикладної математики")
    doc.add_paragraph("Кафедра Системного програмування і спеціалізованих комп’ютерних систем")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Комплексна лабораторна робота")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_page_break()


def add_table_from_matrix(doc, matrix):
    rows_count = len(matrix)
    cols_count = len(matrix[0])
    table = doc.add_table(rows=rows_count, cols=cols_count)



def input_task(doc, data):
    """1. Исходное задание;"""
    doc.add_paragraph("1. Початкове завдання")
    doc.add_picture(data['input'][DEVICE_SCHEME])
    doc.add_page_break()


def task_function_list(doc, data):
    """2. Перечень составленных функций задач;"""
    doc.add_paragraph("2. Список складених функцій задач")
    for i, f in enumerate(data['input'][FNS]):
        doc.add_paragraph("f%i" % (i+1) + ' = ' + f.dump())


def load_balancing(doc, data):
    """3. Таблица перераспределения нагрузки;"""
    doc.add_paragraph("3. Таблиця перерозподілу навантаження")
    add_table_from_matrix(doc, data['input'][LOADS])


def reject_modules_stats(doc, data):
    """4. Статистика отказавших модулей;"""
    pass


def alg_docs(doc, data):
    """5. Детальное описание алгоритма статистической выборки векторов кратности 3 и 4 (1/2 страницы);"""
    pass


def reliability_p_or_q(doc, data):
    """6. Найденный показатель надёжности системы (P или Q);"""
    pass


def modification_proposes(doc, data):
    """7. С учётом п.5 предложения по модификации системы (с обоснованием в плане стоимости, сложности и проч.);"""
    pass


def new_system_structure(doc, data):
    """8. Структура модифицированной системы;"""
    pass


def conclusions(doc, data):
    """10. Вывод о том, что «малой кровью» достигнуто существенное
    (показать на сколько) улучшение показателя надёжности."""
    pass


def system_docs(doc, data):
    load_balancing(doc, data)
    reject_modules_stats(doc, data)
    alg_docs(doc, data)
    reliability_p_or_q(doc, data)


def new_system_docs(doc, data):
    """9. Пункты 2-4 и 6 для модифицированной системы;"""
    system_docs(doc, data)


def document_do_all(doc: Document, data, *funcs):
    for func in funcs:
        func(doc, data)


def dump_report(data, output, pathname, author):
    doc = Document()
    document_do_all(doc,
                    {
                        'author': author,
                        'input': data,
                        'output': output,
                    },
                    title,
                    input_task,
                    task_function_list,
                    system_docs,
                    modification_proposes,
                    new_system_structure,
                    new_system_docs,
                    conclusions)
    doc.save(pathname)
